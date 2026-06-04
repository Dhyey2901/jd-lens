"""Utility helpers — multi-format document text extraction."""
from __future__ import annotations

import io
import re


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF. Works best with text-based PDFs.

    Image-only (scanned) PDFs return an empty string.
    """
    import pypdf

    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    raw = "\n\n".join(p.strip() for p in pages if p.strip())
    cleaned = re.sub(r"\n{3,}", "\n\n", raw)
    return "\n".join(line.rstrip() for line in cleaned.splitlines()).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file, preserving paragraph structure.

    DOCX extraction is lossless for text — no multi-column or OCR issues.
    """
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # Also extract text from tables (common in resumes)
    for table in doc.tables:
        for row in table.rows:
            row_text = "  |  ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    return "\n".join(paragraphs).strip()


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode a plain-text file, normalising whitespace."""
    raw = file_bytes.decode("utf-8", errors="replace")
    cleaned = re.sub(r"\n{3,}", "\n\n", raw)
    return "\n".join(line.rstrip() for line in cleaned.splitlines()).strip()


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    if ext in ("docx",):
        return extract_text_from_docx(file_bytes)
    if ext in ("txt", "md"):
        return extract_text_from_txt(file_bytes)
    raise ValueError(f"Unsupported file type: .{ext}")


SUPPORTED_TYPES = ["pdf", "docx", "txt"]
SUPPORTED_LABEL = "PDF, DOCX, or TXT"
